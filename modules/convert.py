import os
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import pdfplumber
from docx import Document
from openpyxl import Workbook
from fpdf import FPDF

# --- CONFIG PATH ---
BASE = "workspace/convert"
INPUT_DIR = os.path.join(BASE, "input")
OUTPUT_DIR = os.path.join(BASE, "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# --- OCR Setup ---
pytesseract.pytesseract.tesseract_cmd = "tesseract"  # asumsi sudah di PATH
OCR_LANG = "ind+eng"

# ============ UTIL ============
def list_pdfs():
    return [f for f in os.listdir(INPUT_DIR) if f.lower().endswith(".pdf")]

def list_images():
    exts = (".png", ".jpg", ".jpeg")
    return [f for f in os.listdir(INPUT_DIR) if f.lower().endswith(exts)]

def pause_convert():
    input("\nTekan ENTER untuk kembali ke menu convert...")

# ============ CONVERSIONS ============
def pdf_to_images():
    print("=== PDF → IMAGE ===")
    for fname in list_pdfs():
        path = os.path.join(INPUT_DIR, fname)
        pdf = fitz.open(path)
        base = os.path.splitext(fname)[0]
        for i, page in enumerate(pdf, start=1):
            pix = page.get_pixmap(dpi=200)
            out = os.path.join(OUTPUT_DIR, f"{base}_page{i}.jpg")
            pix.save(out)
        pdf.close()
        print("✅", fname)
    print("\nSelesai. Hasil di folder output.")
    pause_convert()

def images_to_pdf():
    print("=== IMAGE → PDF ===")
    images = [f for f in list_images()]
    if not images:
        print("Tidak ada file gambar di input/.")
        return
    pdf = FPDF()
    for fname in images:
        img_path = os.path.join(INPUT_DIR, fname)
        cover = Image.open(img_path)
        w, h = cover.size
        pdf_w, pdf_h = 210, 297
        ratio = min(pdf_w / w, pdf_h / h)
        new_w, new_h = w * ratio, h * ratio
        pdf.add_page()
        pdf.image(img_path, x=(210 - new_w) / 2, y=(297 - new_h) / 2, w=new_w, h=new_h)
    out = os.path.join(OUTPUT_DIR, "merged_images.pdf")
    pdf.output(out)
    print("✅ Berhasil:", out)
    pause_convert()

def pdf_to_word():
    print("=== PDF → WORD (OCR) ===")
    for fname in list_pdfs():
        path = os.path.join(INPUT_DIR, fname)
        base = os.path.splitext(fname)[0]
        doc = Document()
        pdf = fitz.open(path)
        for i, page in enumerate(pdf, start=1):
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(img, lang=OCR_LANG)
            doc.add_paragraph(text)
            doc.add_page_break()
        pdf.close()
        out = os.path.join(OUTPUT_DIR, f"{base}.docx")
        doc.save(out)
        print("✅", fname)
    print("\nSelesai.")
    pause_convert()

def word_to_pdf():
    print("=== WORD → PDF ===")
    from docx import Document
    for fname in os.listdir(INPUT_DIR):
        if not fname.lower().endswith(".docx"):
            continue
        path = os.path.join(INPUT_DIR, fname)
        base = os.path.splitext(fname)[0]
        doc = Document(path)
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Helvetica", size=11)
        for p in doc.paragraphs:
            pdf.multi_cell(0, 8, p.text)
            pdf.ln(2)
        out = os.path.join(OUTPUT_DIR, f"{base}.pdf")
        pdf.output(out)
        print("✅", fname)
    pause_convert()

def pdf_to_excel():
    print("=== PDF → EXCEL (OCR + TABEL) ===")
    for fname in list_pdfs():
        path = os.path.join(INPUT_DIR, fname)
        base = os.path.splitext(fname)[0]
        wb = Workbook()
        ws = wb.active
        ws.title = "Data"
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                table = page.extract_table()
                if table:
                    for row in table:
                        ws.append(row)
                else:
                    img = page.to_image(resolution=200)
                    tmp = "temp_ocr.png"
                    img.save(tmp, format="PNG")
                    text = pytesseract.image_to_string(Image.open(tmp), lang=OCR_LANG)
                    ws.append([text])
                    os.remove(tmp)
        out = os.path.join(OUTPUT_DIR, f"{base}.xlsx")
        wb.save(out)
        print("✅", fname)
    pause_convert()

def excel_to_pdf():
    print("=== EXCEL → PDF ===")
    from openpyxl import load_workbook
    for fname in os.listdir(INPUT_DIR):
        if not fname.lower().endswith(".xlsx"):
            continue
        path = os.path.join(INPUT_DIR, fname)
        wb = load_workbook(path)
        ws = wb.active
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=10)
        for row in ws.iter_rows(values_only=True):
            line = " | ".join(str(x) if x is not None else "" for x in row)
            pdf.multi_cell(0, 6, line)
        out = os.path.join(OUTPUT_DIR, f"{os.path.splitext(fname)[0]}.pdf")
        pdf.output(out)
        print("✅", fname)
    pause_convert()

# ============ MENU ============
def run_convert():
    while True:
        os.system('cls' if os.name == 'nt' else 'clear')
        print("=== CONVERT MENU ===")
        print("1. PDF → Image")
        print("2. Image → PDF")
        print("3. PDF → Word (OCR)")
        print("4. Word → PDF")
        print("5. PDF → Excel (OCR)")
        print("6. Excel → PDF")
        print("0. Kembali ke menu utama")

        ch = input("\nPilih menu: ").strip()
        os.system('cls' if os.name == 'nt' else 'clear')
        if ch == "1":
            pdf_to_images()
        elif ch == "2":
            images_to_pdf()
        elif ch == "3":
            pdf_to_word()
        elif ch == "4":
            word_to_pdf()
        elif ch == "5":
            pdf_to_excel()
        elif ch == "6":
            excel_to_pdf()
        elif ch == "0":
            break
        else:
            print("Pilihan tidak valid.")
            pause_convert()